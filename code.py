import os
import cv2
import pytesseract
import shutil
import pdfplumber
from docx import Document
from docx.shared import RGBColor, Pt
from difflib import unified_diff
from docx import Document as DocxDocument
from docx.oxml import OxmlElement
import fitz  # PyMuPDF
import matplotlib.pyplot as plt
import pandas as pd

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe'


# Function to calculate accuracy
def calculate_accuracy(template_text, target_text, is_template_present):
    template_words = set(template_text.lower().split())
    target_words = set(target_text.lower().split())

    common_words = template_words.intersection(target_words)
    content_matching_ratio = len(common_words) / len(template_words)

    accuracy_content = content_matching_ratio if is_template_present else 0.0
    template_matching = 1.0 if is_template_present else 0.0

    return accuracy_content, template_matching


# Function to extract text from a PDF using pdfplumber
def extract_text_from_pdf(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        text = ""
        for page_number in range(len(pdf.pages)):
            page = pdf.pages[page_number]
            text += page.extract_text()
    return text


# Function to count tables in a PDF using pdfplumber
def count_tables_pdfplumber(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        table_count = sum(len(page.extract_tables()) for page in pdf.pages)
    return table_count


# Function to move a file from source to destination
def move_file(source_path, destination_folder):
    destination_path = os.path.join(destination_folder, os.path.basename(source_path))
    shutil.copy(source_path, destination_path)


# Function to highlight differences and color-code match and mismatch in DOCX
def highlight_diff_in_docx(doc, template_text, target_text):
    diff_result = list(unified_diff(template_text.splitlines(), target_text.splitlines(), lineterm=''))

    # Create a new paragraph
    paragraph = doc.add_paragraph()

    # Iterate through the differences and add text with appropriate color
    for line in diff_result:
        run = paragraph.add_run(line[1:])
        if line.startswith('+'):
            font = run.font
            font.color.rgb = RGBColor(255, 0, 0)  # Set font color to red for mismatch
        elif line.startswith('-'):
            font = run.font
            font.color.rgb = RGBColor(0, 128, 0)  # Set font color to green for match


# Function to save differences to DOCX
def save_diff_to_docx(target_file_name, template_text, target_text, result_folder, accuracy_content, accuracy_template):
    doc = Document()
    doc.add_heading(f'Differences in {target_file_name}', level=1)

    # Highlight differences and color-code match and mismatch
    highlight_diff_in_docx(doc, template_text, target_text)

    # Save accuracy information
    doc.add_paragraph(f'Template Accuracy: {accuracy_template * 100:.2f}%')
    doc.add_paragraph(f'Target Accuracy: {accuracy_content * 100:.2f}%')

    diff_docx_path = os.path.join(result_folder, f'Differences_{target_file_name}.docx')
    doc.save(diff_docx_path)


# Function to plot and save accuracy graph
def plot_accuracy_graph(data, output_path, title, xlabel, ylabel, colors=None):
    df = pd.DataFrame(data, columns=['Target File', 'Match (%)', 'Mismatch (%)'])

    # Plotting
    plt.figure(figsize=(10, 6))
    plt.bar(df['Target File'], df['Match (%)'], color=colors[0], label='Match')
    plt.bar(df['Target File'], df['Mismatch (%)'], color=colors[1], label='Mismatch', bottom=df['Match (%)'])
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.title(title)
    plt.xticks(rotation=45, ha='right')
    plt.legend()
    plt.tight_layout()

    # Save the plot to an image file
    plt.savefig(output_path)
    plt.close()


# Function to create a pie chart
def save_pie_chart(data, labels, title, output_path, colors=None):
    plt.figure(figsize=(6, 6))
    plt.pie(data, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
    plt.title(title)
    plt.savefig(output_path)
    plt.close()

# Function to create a table in the DOCX document
def create_table(doc, data, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'TableGrid'

    # Add header row
    for col, header in enumerate(headers):
        table.cell(0, col).text = header

    # Add data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for col, value in enumerate(row_data):
            row_cells[col].text = str(value)

# Main function for template matching with OCR
def template_matching_with_ocr(template_folder, target_folder, is_template_present, increased_threshold, result_folder,
                               fake_folder):
    if not os.path.exists(result_folder):
        os.makedirs(result_folder)

    if not os.path.exists(fake_folder):
        os.makedirs(fake_folder)

    template_files = [os.path.join(template_folder, file) for file in os.listdir(template_folder) if
                      file.lower().endswith('.pdf')]

    accuracy_data = []  # to store accuracy data for plotting
    content_match_percentages = []  # to store content match percentages for pie chart
    template_match_percentages = []  # to store template match percentages for pie chart

    for target_file_name in os.listdir(target_folder):
        if target_file_name.lower().endswith('.pdf'):
            current_target_path = os.path.join(target_folder, target_file_name)

            template_matched = False

            for template_path in template_files:
                template_name = os.path.basename(template_path)

                # Extract text from the template PDF
                template_text = extract_text_from_pdf(template_path)

                # Extract text from the target PDF
                target_text = extract_text_from_pdf(current_target_path)

                print(f"\nTemplate: {template_name} | Target: {target_file_name}")

                # Compare the extracted text
                accuracy_content, template_matching = calculate_accuracy(template_text, target_text,
                                                                         is_template_present)

                # Display accuracy values
                print(f"Template Accuracy: {template_matching * 100:.2f}%")
                print(f"Content Accuracy: {accuracy_content * 100:.2f}%")

                # Count tables in template and target
                template_table_count = count_tables_pdfplumber(template_path)
                target_table_count = count_tables_pdfplumber(current_target_path)

                # Check if the number of tables matches
                tables_matching = template_table_count == target_table_count

                # Print the results
                template_match_status = "successfully" if template_matching > 0 and tables_matching else "unsuccessfully"
                print(f"\nTemplate {template_name} matched {template_match_status}.")

                # Display content matching status
                content_matching_status = "Match" if accuracy_content == 1.0 else "Different"
                print(f"Content Matching Status: {content_matching_status}")

                # Display more information if template matched
                if template_matching > 0 and accuracy_content == 1.0 and tables_matching:
                    print(f"\nExtracted content from {target_file_name}:")
                    print(target_text)

                    # Move the file to the result_folder if both template and content match
                    move_file(current_target_path, result_folder)
                    template_matched = True
                    break  # Break out of the loop since we found a match

            if not template_matched:
                # If no match found, move the file to the fake_folder
                move_file(current_target_path, fake_folder)
                print(f"Template not matched for {target_file_name}.")

                # Display extracted text for unmatched file
                print(f"\nExtracted content from unmatched file {target_file_name}:")
                print(target_text)

                # Save differences to a Word document and highlight in PDF
                save_diff_to_docx(target_file_name, template_text, target_text, fake_folder,
                                  accuracy_content, template_matching)

            # Append accuracy data for plotting
            accuracy_data.append((target_file_name, template_matching * 100, (1 - accuracy_content) * 100))
            content_match_percentages.append(accuracy_content * 100)
            template_match_percentages.append(template_matching * 100)

    # Create a DOCX document to display the accuracy in tabular format and the graph
    doc = Document()
    doc.add_heading('Comparison Results', level=1)

    # Create a table for accuracy
    accuracy_table_data = [['Target File', 'Template Match (%)', 'Content Mismatch (%)']] + accuracy_data
    create_table(doc, accuracy_table_data, headers=['Target File', 'Template Match (%)', 'Content Mismatch (%)'])

    # Add a new line before the graphs
    doc.add_paragraph()

    # Plot content matching graph and save to DOCX
    content_matching_graph_path = os.path.join(result_folder, 'content_matching_graph.png')
    plot_accuracy_graph(accuracy_data, content_matching_graph_path, 'Content Matching Comparison',
                        'Target File', 'Accuracy (%)', colors=['green', 'red'])
    doc.add_picture(content_matching_graph_path, width=Pt(400), height=Pt(300))  # Adjust the width and height as needed

    # Add a new line before the next graph
    doc.add_paragraph()

    # Plot template matching graph and save to DOCX
    template_matching_graph_path = os.path.join(result_folder, 'template_matching_graph.png')
    plot_accuracy_graph(accuracy_data, template_matching_graph_path, 'Template Matching Comparison',
                        'Target File', 'Accuracy (%)', colors=['blue', 'orange'])
    doc.add_picture(template_matching_graph_path, width=Pt(400), height=Pt(300))  # Adjust the width and height as needed

    # Add a new line before the pie charts
    doc.add_paragraph()

    # Create a pie chart for content matching
    content_pie_chart_path = os.path.join(result_folder, 'content_pie_chart.png')
    save_pie_chart([sum(content_match_percentages), 100 - sum(content_match_percentages)],
                   ['Match', 'Mismatch'], 'Content Matching Status', content_pie_chart_path,
                   colors=['green', 'red'])
    doc.add_paragraph('Content Matching Status:')
    doc.add_picture(content_pie_chart_path, width=Pt(300), height=Pt(300))

    # Add a new line before the next pie chart
    doc.add_paragraph()

    # Create a pie chart for template matching
    template_pie_chart_path = os.path.join(result_folder, 'template_pie_chart.png')
    save_pie_chart([sum(template_match_percentages), 100 - sum(template_match_percentages)],
                   ['Match', 'Mismatch'], 'Template Matching Status', template_pie_chart_path,
                   colors=['blue', 'orange'])
    doc.add_paragraph('Template Matching Status:')
    doc.add_picture(template_pie_chart_path, width=Pt(300), height=Pt(300))

    # Save the DOCX document
    doc.save(os.path.join(result_folder, 'Comparison_Results.docx'))


if _name_ == "_main_":
    template_folder = 'fake'  # Replace with the path to your template folder (contains PDF files)
    target_folder = 'pdf'  # Replace with the path to your target folder (contains PDF files)
    is_template_present = True  # Replace with the ground truth information
    increased_threshold = 0.9  # Adjust the threshold value as needed
    result_folder = 'resullt_folder'  # Replace with the path to your result folder
    fake_folder = 'fake_folder'  # Replace with the path to your fake folder

    # Ensure that the result_folder and fake_folder exist
    os.makedirs(result_folder, exist_ok=True)
    os.makedirs(fake_folder, exist_ok=True)

    template_matching_with_ocr(template_folder, target_folder, is_template_present, increased_threshold, result_folder,
                               fake_folder)
